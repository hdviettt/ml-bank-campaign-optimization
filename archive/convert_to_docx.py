"""
Convert report.md to Google Docs compatible .docx format with embedded images
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path

print("Converting report.md to report.docx...")
print("="*80)

# Create document
doc = Document()

# Configure styles
styles = doc.styles

# Title style
title_style = styles['Title']
title_font = title_style.font
title_font.name = 'Arial'
title_font.size = Pt(24)
title_font.bold = True

# Heading 1
heading1_style = styles['Heading 1']
heading1_font = heading1_style.font
heading1_font.name = 'Arial'
heading1_font.size = Pt(18)
heading1_font.bold = True
heading1_font.color.rgb = RGBColor(0, 0, 139)

# Heading 2
heading2_style = styles['Heading 2']
heading2_font = heading2_style.font
heading2_font.name = 'Arial'
heading2_font.size = Pt(14)
heading2_font.bold = True

# Normal style
normal_style = styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Calibri'
normal_font.size = Pt(11)

# Read the markdown file
with open('report.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Split into lines
lines = content.split('\n')

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    # Skip empty lines at document start
    if not line and i == 0:
        i += 1
        continue

    # Main title (first # heading)
    if line.startswith('# ') and i < 5:
        title_text = line[2:].strip()
        doc.add_heading(title_text, level=0)
        i += 1
        continue

    # Heading 1
    elif line.startswith('## '):
        heading_text = line[3:].strip()
        doc.add_heading(heading_text, level=1)
        i += 1
        continue

    # Heading 2
    elif line.startswith('### '):
        heading_text = line[4:].strip()
        doc.add_heading(heading_text, level=2)
        i += 1
        continue

    # Heading 3
    elif line.startswith('#### '):
        heading_text = line[5:].strip()
        doc.add_heading(heading_text, level=3)
        i += 1
        continue

    # Images
    elif line.startswith('!['):
        # Extract image info: ![alt text](path)
        match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if match:
            alt_text = match.group(1)
            image_path = match.group(2)

            # Check if image exists
            if Path(image_path).exists():
                try:
                    # Add image
                    doc.add_picture(image_path, width=Inches(6.0))
                    # Add caption
                    caption = doc.add_paragraph()
                    caption_run = caption.add_run(alt_text)
                    caption_run.italic = True
                    caption_run.font.size = Pt(10)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"  [OK] Embedded: {image_path}")
                except Exception as e:
                    print(f"  [ERROR] Could not embed {image_path}: {e}")
                    doc.add_paragraph(f"[Image: {alt_text}]")
            else:
                print(f"  [ERROR] Image not found: {image_path}")
                doc.add_paragraph(f"[Image not found: {image_path}]")
        i += 1
        continue

    # Tables (detect markdown tables)
    elif '|' in line and i + 1 < len(lines) and '|' in lines[i+1]:
        # Collect table rows
        table_lines = []
        temp_i = i
        while temp_i < len(lines) and '|' in lines[temp_i]:
            table_lines.append(lines[temp_i])
            temp_i += 1

        # Parse table
        if len(table_lines) >= 2:
            # Extract rows
            rows = []
            for table_line in table_lines:
                cells = [cell.strip() for cell in table_line.split('|')]
                cells = [c for c in cells if c]  # Remove empty
                if cells and not all(c.startswith('-') for c in cells):  # Skip separator line
                    rows.append(cells)

            if rows:
                # Create table
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'

                # Fill table
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # Bold header row
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True

                doc.add_paragraph()  # Add space after table
                print(f"  [OK] Added table with {len(rows)} rows")

        i = temp_i
        continue

    # Bullet lists
    elif line.startswith('- ') or line.startswith('* '):
        bullet_text = line[2:].strip()
        # Handle bold/italic markdown
        bullet_text = re.sub(r'\*\*(.*?)\*\*', r'\1', bullet_text)  # Remove **
        bullet_text = re.sub(r'\*(.*?)\*', r'\1', bullet_text)  # Remove *
        doc.add_paragraph(bullet_text, style='List Bullet')
        i += 1
        continue

    # Numbered lists
    elif re.match(r'^\d+\.', line):
        list_text = re.sub(r'^\d+\.\s*', '', line)
        list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)
        doc.add_paragraph(list_text, style='List Number')
        i += 1
        continue

    # Code blocks
    elif line.startswith('```'):
        # Collect code block
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1
        i += 1  # Skip closing ```

        if code_lines:
            code_text = '\n'.join(code_lines)
            code_para = doc.add_paragraph(code_text)
            code_para.style = 'No Spacing'
            for run in code_para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        continue

    # Horizontal rules
    elif line.startswith('---'):
        doc.add_paragraph()
        i += 1
        continue

    # Block quotes
    elif line.startswith('>'):
        quote_text = line[1:].strip()
        quote_para = doc.add_paragraph(quote_text)
        quote_para.style = 'Intense Quote'
        i += 1
        continue

    # Regular paragraphs
    elif line.strip():
        # Clean up markdown formatting
        paragraph_text = line

        # Handle bold and italic (simple approach - may need refinement)
        # For now, just remove markdown syntax
        paragraph_text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', paragraph_text)  # Bold+italic
        paragraph_text = re.sub(r'\*\*(.*?)\*\*', r'\1', paragraph_text)  # Bold
        paragraph_text = re.sub(r'\*(.*?)\*', r'\1', paragraph_text)  # Italic
        paragraph_text = re.sub(r'`(.*?)`', r'\1', paragraph_text)  # Inline code

        # Add paragraph
        p = doc.add_paragraph(paragraph_text)
        i += 1
        continue

    # Empty line
    else:
        # Only add paragraph break if previous wasn't empty
        if i > 0 and lines[i-1].strip():
            doc.add_paragraph()
        i += 1

# Save document
output_path = 'report.docx'
doc.save(output_path)

print("="*80)
print(f"Successfully created: {output_path}")
print(f"\nNext steps:")
print(f"1. Open {output_path} to verify")
print(f"2. Upload to Google Drive")
print(f"3. Right-click -> Open with -> Google Docs")
print(f"4. It will automatically convert to Google Docs format!")
print("="*80)
